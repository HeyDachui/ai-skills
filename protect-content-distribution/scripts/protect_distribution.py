#!/usr/bin/env python3
"""Add visible official-source guidance and copy IDs to TXT, DOCX, or PDF files."""

from __future__ import annotations

import argparse
import codecs
import datetime as dt
import hashlib
import io
import json
import os
import random
import re
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path


CHAPTER_RE = re.compile(r"^\s*(?:第[〇零一二三四五六七八九十百千万0-9]+[章节卷回部篇]|卷[〇零一二三四五六七八九十百千万0-9]+)")


@dataclass(frozen=True)
class EncodingSpec:
    label: str
    codec: str
    bom: bytes = b""


ENCODINGS = {
    "utf-8": EncodingSpec("utf-8", "utf-8"),
    "utf-8-sig": EncodingSpec("utf-8-sig", "utf-8", codecs.BOM_UTF8),
    "gbk": EncodingSpec("gbk/cp936", "gbk"),
    "gb18030": EncodingSpec("gb18030", "gb18030"),
    "utf-16-le": EncodingSpec("utf-16-le", "utf-16-le", codecs.BOM_UTF16_LE),
    "utf-16-be": EncodingSpec("utf-16-be", "utf-16-be", codecs.BOM_UTF16_BE),
}


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def strict_decode(raw: bytes) -> tuple[str, EncodingSpec]:
    for bom, spec in (
        (codecs.BOM_UTF8, ENCODINGS["utf-8-sig"]),
        (codecs.BOM_UTF16_LE, ENCODINGS["utf-16-le"]),
        (codecs.BOM_UTF16_BE, ENCODINGS["utf-16-be"]),
    ):
        if raw.startswith(bom):
            return raw[len(bom):].decode(spec.codec, errors="strict"), spec
    for name in ("utf-8", "gbk", "gb18030"):
        spec = ENCODINGS[name]
        try:
            return raw.decode(spec.codec, errors="strict"), spec
        except UnicodeDecodeError:
            pass
    raise UnicodeError("unable to decode TXT as UTF-8, GBK/CP936, GB18030, or BOM-marked UTF-16")


def encode_text(text: str, spec: EncodingSpec) -> bytes:
    return spec.bom + text.encode(spec.codec, errors="strict")


def atomic_write(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(dir=path.parent, prefix=path.name + ".", suffix=".tmp", delete=False) as handle:
        temp_path = Path(handle.name)
        handle.write(data)
    try:
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def official_target(args: argparse.Namespace) -> str:
    return args.purchase_url or args.url


def stable_rng(source: Path, copy_id: str, channel: str) -> tuple[random.Random, str]:
    seed_material = source.read_bytes() + b"\0" + copy_id.encode("utf-8") + b"\0" + channel.encode("ascii")
    seed_hash = hashlib.sha256(seed_material).hexdigest()
    return random.Random(int(seed_hash[:16], 16)), seed_hash


def notices(args: argparse.Namespace) -> list[str]:
    target = official_target(args)
    base = args.notice or f"正版提示：本内容由 {args.brand} 正式发布。正版购买与更新：{target}"
    return [
        f"{base}｜副本编号：{args.copy_id}",
        f"正版来源：{args.brand}｜官方入口：{args.url}｜请支持正版｜副本编号：{args.copy_id}",
        f"如本内容来自转载，请通过 {target} 获取正版版本与后续更新｜副本编号：{args.copy_id}",
        f"阅读提示：当前副本编号为 {args.copy_id}。完整正版与修订内容请访问 {target}",
        f"本段之间的提示用于识别非授权转载。正版发布方：{args.brand}｜{args.url}｜{args.copy_id}",
        f"若你通过转载获得本内容，请前往 {target} 支持正版并获取完整更新｜副本 {args.copy_id}",
        f"正版入口不会隐藏：{args.url}｜权利方：{args.brand}｜副本识别码：{args.copy_id}",
    ]


def protect_txt(source: Path, output: Path, args: argparse.Namespace) -> dict[str, object]:
    raw = source.read_bytes()
    text, encoding = strict_decode(raw)
    newline = "\r\n" if "\r\n" in text else "\n"
    had_final = text.endswith(("\n", "\r"))
    lines = text.splitlines()
    templates = notices(args)
    rng, placement_seed = stable_rng(source, args.copy_id, "txt")
    interval = args.txt_interval or ({1: 1200, 2: 600, 3: 300}.get(args.level) or rng.randint(8, 18))
    chapter_interval = args.txt_chapter_interval or {1: 999999, 2: 6, 3: 3, 4: 1}[args.level]
    banner = f"【{templates[0]}】"
    output_lines: list[str] = [banner, ""]
    placements: list[dict[str, object]] = [{"kind": "start", "before_source_line": 1, "template": 1}]
    nonblank_since = 0
    chapter_count = 0
    template_index = 1
    pending_interval = False

    for source_index, line in enumerate(lines, 1):
        chapter_hit = bool(CHAPTER_RE.match(line))
        if chapter_hit:
            chapter_count += 1
        chapter_due = chapter_hit and chapter_count > 1 and chapter_count % chapter_interval == 0
        interval_due = pending_interval and (not line.strip() or args.level == 4)
        if chapter_due or interval_due:
            chosen_index = rng.randrange(len(templates)) if args.level == 4 else template_index % len(templates)
            chosen = templates[chosen_index]
            output_lines.extend([f"【{chosen}】", ""])
            placements.append({
                "kind": "chapter" if chapter_due else "interval",
                "before_source_line": source_index,
                "template": chosen_index + 1,
            })
            template_index += 1
            nonblank_since = 0
            pending_interval = False
            if args.level == 4 and args.txt_interval is None:
                interval = rng.randint(8, 18)
        output_lines.append(line)
        if line.strip():
            nonblank_since += 1
            if nonblank_since >= interval:
                pending_interval = True

    if output_lines and output_lines[-1] != "":
        output_lines.append("")
    output_lines.append(f"【{templates[template_index % len(templates)]}】")
    placements.append({"kind": "end", "after_source_line": len(lines), "template": template_index % len(templates) + 1})
    output_text = newline.join(output_lines) + (newline if had_final else "")
    output_bytes = encode_text(output_text, encoding)
    atomic_write(output, output_bytes)
    verified_text, verified_encoding = strict_decode(output.read_bytes())
    return {
        "format": "txt",
        "encoding": encoding.label,
        "bom": bool(encoding.bom),
        "newline": "CRLF" if newline == "\r\n" else "LF",
        "source_line_count": len(lines),
        "output_line_count": len(verified_text.splitlines()),
        "source_chapter_count": sum(bool(CHAPTER_RE.match(line)) for line in lines),
        "output_chapter_count": sum(bool(CHAPTER_RE.match(line)) for line in verified_text.splitlines()),
        "placement_strategy": "deterministic_random_body_interruption" if args.level == 4 else "fixed_structural_intervals",
        "placement_seed_sha256": placement_seed if args.level == 4 else None,
        "placements": placements,
        "verification": {
            "strict_reopen": True,
            "encoding_preserved": verified_encoding.label == encoding.label,
            "replacement_chars": verified_text.count("\ufffd"),
            "official_url_present": args.url in verified_text,
            "copy_id_present": args.copy_id in verified_text,
        },
    }


def add_hyperlink(paragraph, text: str, url: str):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "496B8A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_docx_watermark(paragraph, text: str, level: int) -> None:
    from lxml import etree

    fill = {2: "E5E7EB", 3: "D1D5DB", 4: "B8C0CC"}.get(level, "E5E7EB")
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    v_ns = "urn:schemas-microsoft-com:vml"
    o_ns = "urn:schemas-microsoft-com:office:office"
    run = etree.Element(f"{{{w_ns}}}r")
    pict = etree.SubElement(run, f"{{{w_ns}}}pict")
    shape = etree.SubElement(
        pict,
        f"{{{v_ns}}}shape",
        {
            "id": "OfficialSourceWatermark",
            f"{{{o_ns}}}spid": "_x0000_s2049",
            "type": "#_x0000_t136",
            "style": (
                "position:absolute;margin-left:0;margin-top:0;width:470pt;height:110pt;"
                "rotation:315;z-index:-251654144;mso-position-horizontal:center;"
                "mso-position-horizontal-relative:margin;mso-position-vertical:center;"
                "mso-position-vertical-relative:margin;mso-wrap-edited:f"
            ),
            "fillcolor": f"#{fill}",
            "stroked": "f",
        },
    )
    etree.SubElement(shape, f"{{{v_ns}}}fill", {"opacity": {2: "0.10", 3: "0.16", 4: "0.22"}.get(level, "0.10")})
    etree.SubElement(
        shape,
        f"{{{v_ns}}}textpath",
        {"style": 'font-family:"Microsoft YaHei";font-size:1pt', "string": text},
    )
    etree.SubElement(shape, f"{{{v_ns}}}path", {"textpathok": "t"})
    paragraph._p.append(run)


def set_docx_read_only(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings._element
    for existing in settings.findall(qn("w:documentProtection")):
        settings.remove(existing)
    protection = OxmlElement("w:documentProtection")
    protection.set(qn("w:edit"), "readOnly")
    protection.set(qn("w:enforcement"), "1")
    settings.append(protection)


def protect_docx(source: Path, output: Path, args: argparse.Namespace) -> dict[str, object]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document(source)
    original_body_paragraphs = list(document.paragraphs)
    target = official_target(args)
    footer_text = f"正版来源：{args.brand}｜副本编号：{args.copy_id}｜"
    watermark_text = f"正版来源 {args.brand} · {args.copy_id}"
    section_records: list[dict[str, object]] = []

    for section_index, section in enumerate(document.sections, 1):
        section.footer.is_linked_to_previous = False
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)
        add_hyperlink(paragraph, target, target)
        if args.level >= 2:
            section.header.is_linked_to_previous = False
            watermark_paragraph = section.header.paragraphs[0]
            add_docx_watermark(watermark_paragraph, watermark_text, args.level)
        section_records.append({"section": section_index, "footer": True, "watermark": args.level >= 2})

    if args.level >= 3 and document.paragraphs:
        first = document.paragraphs[0]
        notice = first.insert_paragraph_before(f"正版说明｜{args.brand}｜正版入口：{target}｜副本编号：{args.copy_id}")
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in notice.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(73, 107, 138)

    body_notice_records: list[dict[str, object]] = []
    placement_seed = None
    if args.level == 4 and original_body_paragraphs:
        rng, placement_seed = stable_rng(source, args.copy_id, "docx")
        templates = notices(args)
        next_gap = rng.randint(8, 18)
        nonblank_count = 0
        for source_paragraph, paragraph in enumerate(original_body_paragraphs, 1):
            if not paragraph.text.strip():
                continue
            nonblank_count += 1
            if nonblank_count < next_gap:
                continue
            template_index = rng.randrange(len(templates))
            inserted = paragraph.insert_paragraph_before(f"【{templates[template_index]}】")
            inserted.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in inserted.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(73, 107, 138)
            body_notice_records.append({"before_source_paragraph": source_paragraph, "template": template_index + 1})
            nonblank_count = 0
            next_gap = rng.randint(8, 18)

    properties = document.core_properties
    properties.subject = f"Official source: {args.url}"
    properties.comments = f"Official edition guidance; copy ID {args.copy_id}"
    properties.keywords = f"official edition, copyright, {args.copy_id}"
    properties.last_modified_by = "protect-content-distribution"
    if args.docx_read_only:
        set_docx_read_only(document)
    document.save(output)

    with zipfile.ZipFile(output) as archive:
        xml_text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml") or name.endswith(".rels")
        )
    return {
        "format": "docx",
        "sections": section_records,
        "placements": {
            "footer_every_section": True,
            "watermark_every_section": args.level >= 2,
            "front_notice": args.level >= 3,
            "read_only_flag": bool(args.docx_read_only),
            "body_notice_count": len(body_notice_records),
            "body_notices": body_notice_records,
            "placement_strategy": "deterministic_random_body_interruption" if args.level == 4 else None,
            "placement_seed_sha256": placement_seed,
        },
        "verification": {
            "zip_reopen": True,
            "official_url_present": args.url in xml_text or target in xml_text,
            "copy_id_present": args.copy_id in xml_text,
            "watermark_present": watermark_text in xml_text if args.level >= 2 else None,
            "render_required": True,
        },
    }


def register_pdf_font(font_path: Path | None) -> str:
    from reportlab.pdfbase import pdfmetrics
    if font_path:
        from reportlab.pdfbase.ttfonts import TTFont
        name = "ProtectionFont"
        pdfmetrics.registerFont(TTFont(name, str(font_path)))
        return name
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(name))
    except KeyError:
        pass
    return name


def add_pdf_link(writer, page, url: str, rect: tuple[float, float, float, float]) -> None:
    from pypdf.generic import ArrayObject, DictionaryObject, FloatObject, NameObject, TextStringObject

    annotation = DictionaryObject({
        NameObject("/Type"): NameObject("/Annot"),
        NameObject("/Subtype"): NameObject("/Link"),
        NameObject("/Rect"): ArrayObject([FloatObject(value) for value in rect]),
        NameObject("/Border"): ArrayObject([FloatObject(0), FloatObject(0), FloatObject(0)]),
        NameObject("/A"): DictionaryObject({
            NameObject("/S"): NameObject("/URI"),
            NameObject("/URI"): TextStringObject(url),
        }),
    })
    reference = writer._add_object(annotation)
    annotations = page.get("/Annots")
    if annotations is None:
        page[NameObject("/Annots")] = ArrayObject([reference])
    else:
        annotations.get_object().append(reference)


def protect_pdf(source: Path, output: Path, args: argparse.Namespace) -> dict[str, object]:
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(source)
    if reader.is_encrypted:
        raise ValueError("encrypted PDFs are not supported; provide an authorized decrypted source copy")
    writer = PdfWriter()
    font_name = register_pdf_font(args.font)
    target = official_target(args)
    footer = f"正版来源：{args.brand}  |  {target}  |  副本：{args.copy_id}"
    watermark = f"正版来源 {args.brand} · {args.copy_id}"
    opacity = {1: 0.0, 2: 0.08, 3: 0.13, 4: 0.18}[args.level]
    page_records: list[dict[str, object]] = []
    rng, placement_seed = stable_rng(source, args.copy_id, "pdf")
    next_band_page = rng.randint(1, min(3, max(1, len(reader.pages)))) if args.level == 4 else None
    templates = notices(args)

    for page_index, source_page in enumerate(reader.pages, 1):
        if getattr(source_page, "rotation", 0):
            source_page.transfer_rotation_to_content()
        width = float(source_page.mediabox.width)
        height = float(source_page.mediabox.height)
        buffer = io.BytesIO()
        overlay = canvas.Canvas(buffer, pagesize=(width, height))
        from reportlab.pdfbase import pdfmetrics
        footer_size = min(7.5, max(5.5, (width * 0.88) / max(1, pdfmetrics.stringWidth(footer, font_name, 1))))
        overlay.setFont(font_name, footer_size)
        overlay.setFillColorRGB(0.25, 0.32, 0.40)
        overlay.drawCentredString(width / 2, 14, footer)
        if args.level >= 2:
            overlay.saveState()
            if hasattr(overlay, "setFillAlpha"):
                overlay.setFillAlpha(opacity)
            overlay.setFillColorRGB(0.35, 0.40, 0.46)
            overlay.translate(width / 2, height / 2)
            overlay.rotate(35)
            desired_size = {2: 26, 3: 30, 4: 34}.get(args.level, 26)
            watermark_size = min(desired_size, (width * 0.78) / max(1, pdfmetrics.stringWidth(watermark, font_name, 1)))
            overlay.setFont(font_name, watermark_size)
            overlay.drawCentredString(0, 0, watermark)
            overlay.restoreState()
        body_band = args.level == 4 and page_index == next_band_page
        if body_band:
            band_text = templates[rng.randrange(len(templates))]
            band_y = height * rng.uniform(0.34, 0.58)
            overlay.saveState()
            if hasattr(overlay, "setFillAlpha"):
                overlay.setFillAlpha(0.90)
            overlay.setFillColorRGB(0.93, 0.95, 0.97)
            overlay.rect(0, band_y - 22, width, 44, fill=1, stroke=0)
            if hasattr(overlay, "setFillAlpha"):
                overlay.setFillAlpha(1.0)
            overlay.setFillColorRGB(0.18, 0.25, 0.34)
            band_size = min(9.0, max(6.0, (width * 0.90) / max(1, pdfmetrics.stringWidth(band_text, font_name, 1))))
            overlay.setFont(font_name, band_size)
            overlay.drawCentredString(width / 2, band_y - 3, band_text)
            overlay.restoreState()
            next_band_page = page_index + rng.randint(2, 4)
        overlay.save()
        buffer.seek(0)
        overlay_page = PdfReader(buffer).pages[0]
        source_page.merge_page(overlay_page)
        writer.add_page(source_page)
        output_page = writer.pages[-1]
        add_pdf_link(writer, output_page, target, (max(0, width * 0.18), 4, min(width, width * 0.82), 24))
        page_records.append({"page": page_index, "footer": True, "watermark": args.level >= 2, "body_band": body_band, "link": True})

    metadata = dict(reader.metadata or {})
    metadata.update({
        "/Subject": f"Official source: {args.url}",
        "/Keywords": f"official edition, copyright, {args.copy_id}",
        "/Creator": "protect-content-distribution",
    })
    writer.add_metadata({str(key): str(value) for key, value in metadata.items() if value is not None})
    with tempfile.NamedTemporaryFile(dir=output.parent, prefix=output.name + ".", suffix=".tmp", delete=False) as handle:
        temp_path = Path(handle.name)
        writer.write(handle)
    try:
        os.replace(temp_path, output)
    finally:
        if temp_path.exists():
            temp_path.unlink()

    verified = PdfReader(output)
    extracted = "\n".join(page.extract_text() or "" for page in verified.pages)
    link_count = sum(len(page.get("/Annots", [])) for page in verified.pages)
    return {
        "format": "pdf",
        "page_count": len(verified.pages),
        "placements": page_records,
        "placement_strategy": "deterministic_random_body_interruption" if args.level == 4 else "page_furniture",
        "placement_seed_sha256": placement_seed if args.level == 4 else None,
        "verification": {
            "pdf_reopen": True,
            "page_count_preserved": len(verified.pages) == len(reader.pages),
            "official_url_present_in_text": target in extracted,
            "copy_id_present_in_text": args.copy_id in extracted,
            "link_annotation_count": link_count,
            "render_required": True,
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--brand", required=True)
    parser.add_argument("--url", required=True)
    parser.add_argument("--purchase-url")
    parser.add_argument("--copy-id", required=True)
    parser.add_argument("--notice")
    parser.add_argument("--level", type=int, choices=(1, 2, 3, 4), default=2)
    parser.add_argument("--txt-interval", type=int)
    parser.add_argument("--txt-chapter-interval", type=int)
    parser.add_argument("--docx-read-only", action="store_true")
    parser.add_argument("--font", type=Path)
    parser.add_argument("--manifest", type=Path)
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()

    if args.input.resolve() == args.output.resolve():
        parser.error("output must not overwrite input")
    if args.output.exists() and not args.force:
        parser.error("output exists; use --force to replace the protected copy")
    if args.input.suffix.lower() != args.output.suffix.lower():
        parser.error("input and output must use the same file extension")
    if args.input.suffix.lower() not in {".txt", ".docx", ".pdf"}:
        parser.error("supported formats: .txt, .docx, .pdf")
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._-]{2,63}", args.copy_id):
        parser.error("copy ID must be 3-64 characters using letters, digits, dot, underscore, or hyphen")
    if not args.url.lower().startswith(("https://", "http://")):
        parser.error("--url must start with http:// or https://")
    if args.purchase_url and not args.purchase_url.lower().startswith(("https://", "http://")):
        parser.error("--purchase-url must start with http:// or https://")
    minimum_txt_interval = 5 if args.level == 4 else 20
    if args.txt_interval is not None and args.txt_interval < minimum_txt_interval:
        parser.error(f"--txt-interval must be at least {minimum_txt_interval} nonblank lines for level {args.level}")
    if args.txt_chapter_interval is not None and args.txt_chapter_interval < 1:
        parser.error("--txt-chapter-interval must be at least 1")

    manifest_path = args.manifest or args.output.with_suffix(args.output.suffix + ".protection.json")
    if manifest_path.exists() and not args.force:
        parser.error("manifest exists; use --force to replace the protection manifest")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    source_raw = args.input.read_bytes()
    source_hash = sha256_bytes(source_raw)
    if args.input.suffix.lower() == ".txt":
        details = protect_txt(args.input, args.output, args)
    elif args.input.suffix.lower() == ".docx":
        details = protect_docx(args.input, args.output, args)
    else:
        details = protect_pdf(args.input, args.output, args)

    output_hash = sha256_bytes(args.output.read_bytes())
    source_unchanged = sha256_bytes(args.input.read_bytes()) == source_hash
    manifest = {
        "schema": "protect-content-distribution/v1",
        "created_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "input": str(args.input.resolve()),
        "output": str(args.output.resolve()),
        "input_sha256": source_hash,
        "output_sha256": output_hash,
        "source_unchanged": source_unchanged,
        "brand": args.brand,
        "official_url": args.url,
        "purchase_url": args.purchase_url,
        "copy_id": args.copy_id,
        "protection_level": args.level,
        "details": details,
    }
    atomic_write(manifest_path, (json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8"))
    print(json.dumps({
        "format": details["format"],
        "output": str(args.output.resolve()),
        "manifest": str(manifest_path.resolve()),
        "source_unchanged": source_unchanged,
        "protection_level": args.level,
        "copy_id": args.copy_id,
        "verification": details["verification"],
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
